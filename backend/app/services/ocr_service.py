"""
OCR 服務層
負責 PDF 文件的 OCR 處理
"""

import os
import glob
import sqlite3
import re
import numpy as np
import threading
from pathlib import Path
from typing import List, Dict
import fitz
import io
from PIL import Image
import httpx

from ..paths import get_database_path, get_factory_path


class OCRService:
    """OCR 服務類別"""

    def __init__(self):
        """初始化 OCR 服務"""
        self.ocr = None  # 延遲初始化，避免啟動時載入
        self.db_path = get_database_path()
        self.factory_path = get_factory_path()
        self.is_scanning = False
        self.cancel_requested = False
        self.scan_progress = {"current": 0, "total": 0, "current_file": ""}

    def get_progress(self) -> Dict:
        """取得目前背景掃描進度"""
        return {
            "is_scanning": self.is_scanning,
            "progress": self.scan_progress
        }

    def cancel_scan(self):
        """請求終止背景掃描"""
        if self.is_scanning:
            self.cancel_requested = True

    def start_background_scan(self, filepaths: List[str]) -> bool:
        """開始在背景執行緒中掃描多個檔案"""
        if self.is_scanning:
            return False
            
        self.is_scanning = True
        self.cancel_requested = False
        self.scan_progress = {
            "current": 0,
            "total": len(filepaths),
            "current_file": ""
        }
        
        def run_scan():
            try:
                for i, filepath in enumerate(filepaths):
                    if self.cancel_requested:
                        break
                        
                    p = Path(filepath)
                    self.scan_progress["current"] = i + 1
                    self.scan_progress["current_file"] = p.name
                    try:
                        self.scan_single_file(filepath)
                      # 這裡手動擷取任務已被強制中斷異常，使背景掃描也能順利退出
                    except Exception as e:
                        if "任務已被強制中斷" in str(e):
                            break
                        print(f"Error scanning {filepath}: {e}")
            finally:
                self.is_scanning = False
                self.cancel_requested = False
                
        thread = threading.Thread(target=run_scan)
        thread.daemon = True
        thread.start()
        
        return True

    def ocr_pdf(self, pdf_path: str) -> List[Dict]:
        """
        將 PDF 的每一頁轉為圖片後，發送至容器端 OCR Worker 服務進行辨識

        Args:
            pdf_path: PDF 檔案路徑

        Returns:
            list of dict, 每個 dict 包含 page_number 和 content
        """
        import datetime

        # 使用 PyMuPDF 將 PDF 各頁轉為 PNG byte 流，免除本機端額外依賴
        images_data = []
        doc = fitz.open(pdf_path)
        for page in doc:
            pix = page.get_pixmap(dpi=150)
            images_data.append(pix.tobytes("png"))
        doc.close()

        page_results = []
        # 設定超時為 60 秒，以防 OCR 運算在複雜圖片上耗時較長
        with httpx.Client(timeout=60.0) as client:
            for page_num, img_bytes in enumerate(images_data, start=1):
                # 檢查是否已達預約掃描結束時間，若是則自動中斷
                if hasattr(self, 'scheduled_end_time') and self.scheduled_end_time:
                    end_time = self.scheduled_end_time
                    if isinstance(end_time, str):
                        try:
                            end_time = datetime.datetime.fromisoformat(end_time.replace(" ", "T"))
                        except Exception:
                            pass
                    if isinstance(end_time, datetime.datetime) and datetime.datetime.now() >= end_time:
                        self.cancel_requested = True

                if hasattr(self, 'cancel_requested') and self.cancel_requested:
                    raise Exception("任務已被強制中斷")

                page_text = ""
                try:
                    files = {"file": ("page.png", img_bytes, "image/png")}
                    # 呼叫運行在容器內的 OCR 微服務 (預設埠為 5000)
                    response = client.post("http://localhost:5000/predict", files=files)
                    if response.status_code == 200:
                        res_json = response.json()
                        page_text = "\n".join(res_json.get("rec_texts", []))
                    else:
                        print(f"[OCR] Worker 錯誤 (狀態碼 {response.status_code}): {response.text}")
                except Exception as e:
                    print(f"[OCR] 無法連線至 Docker OCR 服務: {e}")
                    # 在連線失敗時拋出明確異常以提示使用者啟動 Docker
                    raise Exception(f"無法連線至 Docker OCR 服務，請確認容器已啟動 (錯誤: {e})")

                page_results.append({"page_number": page_num, "content": page_text})

        return page_results

    def extract_docx(self, docx_path: str) -> List[Dict]:
        """
        提取 Word (.docx) 文件的文字內容，整份文件做為第 1 頁寫入。
        """
        import docx
        import datetime

        # 檢查超時與中斷
        if hasattr(self, 'scheduled_end_time') and self.scheduled_end_time:
            end_time = self.scheduled_end_time
            if isinstance(end_time, str):
                try:
                    end_time = datetime.datetime.fromisoformat(end_time.replace(" ", "T"))
                except Exception:
                    pass
            if isinstance(end_time, datetime.datetime) and datetime.datetime.now() >= end_time:
                self.cancel_requested = True

        if hasattr(self, 'cancel_requested') and self.cancel_requested:
            raise Exception("任務已被強制中斷")

        doc = docx.Document(docx_path)
        
        paragraphs = []
        # 1. 取得段落文字
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                paragraphs.append(text)
                
        # 2. 取得表格文字
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    # 去重相鄰重複儲存格 (因為合併儲存格會重複讀取)
                    unique_row_text = []
                    for cell_t in row_text:
                        if not unique_row_text or cell_t != unique_row_text[-1]:
                            unique_row_text.append(cell_t)
                    paragraphs.append(" | ".join(unique_row_text))
                    
        full_text = "\n".join(paragraphs)
        return [{"page_number": 1, "content": full_text}]

    def extract_xlsx(self, xlsx_path: str) -> List[Dict]:
        """
        提取 Excel (.xlsx) 文件的文字內容，所有 Sheets 的內容合併做為第 1 頁寫入。
        """
        import openpyxl
        import datetime

        # 檢查超時與中斷
        if hasattr(self, 'scheduled_end_time') and self.scheduled_end_time:
            end_time = self.scheduled_end_time
            if isinstance(end_time, str):
                try:
                    end_time = datetime.datetime.fromisoformat(end_time.replace(" ", "T"))
                except Exception:
                    pass
            if isinstance(end_time, datetime.datetime) and datetime.datetime.now() >= end_time:
                self.cancel_requested = True

        if hasattr(self, 'cancel_requested') and self.cancel_requested:
            raise Exception("任務已被強制中斷")

        wb = openpyxl.load_workbook(xlsx_path, data_only=True, read_only=True)
        try:
            sheets_text = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_lines = []
                for row in sheet.iter_rows(values_only=True):
                    row_text = [str(val).strip() for val in row if val is not None and str(val).strip()]
                    if row_text:
                        sheet_lines.append(" ".join(row_text))
                
                if sheet_lines:
                    sheets_text.append(f"--- 工作表: {sheet_name} ---\n" + "\n".join(sheet_lines))
            
            full_text = "\n\n".join(sheets_text)
            return [{"page_number": 1, "content": full_text}]
        finally:
            wb.close()

    def is_file_in_db(self, filename: str) -> bool:
        """
        檢查檔名是否已存在資料庫

        Args:
            filename: 檔案名稱

        Returns:
            是否存在
        """
        conn = sqlite3.connect(str(self.db_path))
        cursor = conn.cursor()

        try:
            cursor.execute(
                """
                SELECT COUNT(*) FROM documents WHERE filename = ?
            """,
                (filename,),
            )

            count = cursor.fetchone()[0]
            return count > 0
        finally:
            conn.close()

    def save_to_db(self, filename: str, filepath: str, page_results: List[Dict]) -> int:
        """
        將 OCR 或文字提取結果存入資料庫

        Args:
            filename: 檔案名稱
            filepath: 檔案路徑
            page_results: 文字提取結果列表

        Returns:
            文件 ID
        """
        conn = sqlite3.connect(str(self.db_path))
        cursor = conn.cursor()

        try:
            # 嘗試從檔名中萃取出 YYYYMMDD 的日期 (例如: 20250207)
            # 尋找前後帶有連字號的 8 位數字: -20250207-
            doc_date = None
            date_match = re.search(r'-(20[2-9]\d{5})-?', filename)
            if date_match:
                raw_date = date_match.group(1)
                # 轉成 YYYY-MM-DD 格式
                doc_date = f"{raw_date[:4]}-{raw_date[4:6]}-{raw_date[6:8]}"
        
            # 1. 插入 documents 表
            cursor.execute(
                """
                INSERT INTO documents (filename, filepath, doc_date)
                VALUES (?, ?, ?)
            """,
                (filename, filepath, doc_date),
            )

            doc_id = cursor.lastrowid

            # 2. 為每一頁插入 doc_fts 表 (全文檢索)
            for page_data in page_results:
                cursor.execute(
                    """
                    INSERT INTO doc_fts (doc_id, content, page_number)
                    VALUES (?, ?, ?)
                """,
                    (doc_id, page_data["content"], page_data["page_number"]),
                )

            conn.commit()
            return doc_id

        except Exception as e:
            conn.rollback()
            raise Exception(f"存入資料庫失敗: {str(e)}")
        finally:
            conn.close()

    def check_new_files(self, folder_path: str = None) -> Dict:
        """
        檢查指定資料夾中有哪些新的檔案需要處理 (支援 PDF, Docx, Xlsx)
        (不執行 OCR，只返回檔案資訊)

        Returns:
            新檔案的數量和詳細資訊
        """
        target_path = Path(folder_path) if folder_path else self.factory_path
        
        if not target_path.exists() or not target_path.is_dir():
            raise Exception(f"資料夾不存在或無效: {target_path}")

        # 取得所有支援的檔案
        supported_extensions = {".pdf", ".docx", ".xlsx"}
        supported_files = [
            p for p in target_path.iterdir()
            if p.is_file() and p.suffix.lower() in supported_extensions
        ]

        new_files = []
        for file_path in supported_files:
            filename = file_path.name
            filepath = str(file_path.absolute())

            # 檢查是否已在資料庫
            if not self.is_file_in_db(filename):
                file_size = file_path.stat().st_size
                new_files.append(
                    {"filename": filename, "filepath": filepath, "size": file_size}
                )

        return {
            "new_files_count": len(new_files),
            "new_files": new_files,
            "total_files": len(supported_files),
        }

    def scan_factory_folder(self) -> Dict:
        """
        掃描 factory 資料夾下的所有 PDF, Docx, Xlsx 檔案，
        若檔名不在資料庫則進行處理並存入資料庫

        Returns:
            處理結果統計
        """
        if not self.factory_path.exists():
            raise Exception(f"資料夾不存在: {self.factory_path}")

        # 取得所有支援的檔案
        supported_extensions = {".pdf", ".docx", ".xlsx"}
        supported_files = [
            p for p in self.factory_path.iterdir()
            if p.is_file() and p.suffix.lower() in supported_extensions
        ]

        if not supported_files:
            return {
                "total": 0,
                "processed": 0,
                "skipped": 0,
                "failed": 0,
                "details": [],
            }

        processed_count = 0
        skipped_count = 0
        failed_count = 0
        details = []

        for file_path in supported_files:
            filename = file_path.name
            filepath = str(file_path.absolute())

            # 檢查是否已在資料庫
            if self.is_file_in_db(filename):
                skipped_count += 1
                details.append(
                    {
                        "filename": filename,
                        "status": "skipped",
                        "message": "已存在資料庫",
                    }
                )
                continue

            # 進行 OCR / 文字提取
            try:
                ext = file_path.suffix.lower()
                if ext == ".pdf":
                    page_results = self.ocr_pdf(str(file_path))
                    msg = f"成功處理 {len(page_results)} 頁"
                elif ext == ".docx":
                    page_results = self.extract_docx(str(file_path))
                    msg = "成功提取 Word 內文"
                elif ext == ".xlsx":
                    page_results = self.extract_xlsx(str(file_path))
                    msg = "成功提取 Excel 內文"
                else:
                    raise Exception(f"不支援的檔案格式: {ext}")

                doc_id = self.save_to_db(filename, filepath, page_results)
                processed_count += 1
                details.append(
                    {
                        "filename": filename,
                        "status": "success",
                        "message": msg,
                        "doc_id": doc_id,
                        "pages": len(page_results),
                    }
                )
            except Exception as e:
                failed_count += 1
                details.append(
                    {"filename": filename, "status": "failed", "message": str(e)}
                )

        return {
            "total": len(supported_files),
            "processed": processed_count,
            "skipped": skipped_count,
            "failed": failed_count,
            "details": details,
        }

    def scan_single_file(self, filepath: str) -> Dict:
        """
        掃描單一 PDF, Docx, Xlsx 檔案
        
        Args:
            filepath: 檔案路徑
            
        Returns:
            處理結果
        """
        file_path = Path(filepath)
        if not file_path.exists() or not file_path.is_file():
            raise Exception(f"檔案不存在: {filepath}")
            
        filename = file_path.name
        
        # 檢查是否已在資料庫
        if self.is_file_in_db(filename):
            return {
                "success": False,
                "filename": filename,
                "status": "skipped",
                "message": "已存在資料庫"
            }
            
        # 進行 OCR / 文字提取
        try:
            ext = file_path.suffix.lower()
            if ext == ".pdf":
                page_results = self.ocr_pdf(str(file_path))
                msg = f"成功處理 {len(page_results)} 頁"
            elif ext == ".docx":
                page_results = self.extract_docx(str(file_path))
                msg = "成功提取 Word 內文"
            elif ext == ".xlsx":
                page_results = self.extract_xlsx(str(file_path))
                msg = "成功提取 Excel 內文"
            else:
                raise Exception(f"不支援的檔案格式: {ext}")

            doc_id = self.save_to_db(filename, filepath, page_results)
            return {
                "success": True,
                "filename": filename,
                "status": "success",
                "message": msg,
                "doc_id": doc_id,
                "pages": len(page_results)
            }
        except Exception as e:
            return {
                "success": False,
                "filename": filename,
                "status": "failed",
                "message": str(e)
            }

# 建立全域單一實例
ocr_service = OCRService()

